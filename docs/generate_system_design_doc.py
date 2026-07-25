# -*- coding: utf-8 -*-
"""JoyForOld 系统设计说明书 — 扩充技术深度版（竞赛软工文风）"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOCS = Path(__file__).resolve().parent
OUTS = [
    DOCS / "JoyForOld-System-Design-Specification.docx",
    DOCS / "JoyForOld-System-Design-Specification-v2.docx",
]


def font(run, size=12, bold=False, cn="宋体", en="Times New Roman", color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = en
    if color:
        run.font.color.rgb = color
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), cn)


def P(doc, text, size=12, bold=False, align="justify", indent=True, before=0, after=6, cn="宋体"):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.5
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf.first_line_indent = Cm(0.74) if indent and align == "justify" else Cm(0)
    r = p.add_run(text)
    font(r, size=size, bold=bold, cn=cn)
    return p


def H(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for r in p.runs:
        font(r, size={1: 16, 2: 14, 3: 12}[level], bold=True, cn="黑体")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)


def num(doc, items):
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"（{i}）{t}")
        font(r, size=12)


def cell(c, text, bold=False, center=False, size=10.5):
    c.text = ""
    p = c.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    font(r, size=size, bold=bold)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell(t.rows[0].cells[i], h, bold=True, center=True)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            cell(t.rows[ri + 1].cells[ci], str(v))
    doc.add_paragraph()


def pre(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(line)
        font(r, size=9.5, cn="宋体", en="Courier New")


def cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    P(doc, "第21届湖南省大学生计算机程序设计竞赛", size=16, bold=True, align="center", indent=False, cn="黑体")
    P(doc, "——应用开发类竞赛（2025）", size=14, bold=True, align="center", indent=False, cn="黑体", after=18)
    for _ in range(2):
        doc.add_paragraph()
    P(doc, "系  统  设  计  说  明  书", size=28, bold=True, align="center", indent=False, cn="黑体", after=28)
    for _ in range(2):
        doc.add_paragraph()
    for k, v in [
        ("作品名称：", "JoyForOld（面向老年用户的手机语音操控助手）"),
        ("作品类别：", "移动终端开发类"),
        ("作　　者：", "（填写队员姓名）"),
        ("指导老师：", "（填写指导教师）"),
        ("单　　位：", "（填写学校名称）"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(k)
        font(r1, size=14, bold=True, cn="黑体")
        r2 = p.add_run(v)
        font(r2, size=14)
    for _ in range(4):
        doc.add_paragraph()
    P(doc, "二零二五年   月   日", size=14, align="center", indent=False)
    doc.add_page_break()


def toc(doc):
    P(doc, "目    录", size=16, bold=True, align="center", indent=False, cn="黑体", after=12)
    P(doc, "用 Word 打开后，在下方目录域上右键“更新域”生成目录。打印胶装要求见页脚。",
      size=10.5, indent=False, after=12)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run._r.append(run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"}))
    run2 = paragraph.add_run()
    instr = run2._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = r' TOC \o "1-3" \h \z \u '
    run2._r.append(instr)
    run3 = paragraph.add_run()
    run3._r.append(run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"}))
    run4 = paragraph.add_run("（目录域，请更新）")
    font(run4, size=12, color=RGBColor(0x66, 0x66, 0x66))
    run5 = paragraph.add_run()
    run5._r.append(run5._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"}))
    doc.add_page_break()


def body(doc):
    # ========== 1 ==========
    H(doc, "1 引言", 1)
    H(doc, "1.1 编写目的", 2)
    P(doc, "本文档给出 JoyForOld 的系统设计方案，包括总体结构、关键算法与数据流、模块接口、"
      "存储方案、出错处理与安全机制，作为实现、测试和竞赛评审的依据。")

    H(doc, "1.2 项目背景与要解决的问题", 2)
    P(doc, "老年用户使用手机时的主要困难不在“不会开机”，而在“不会在具体 App 里连续操作”。"
      "打电话、发微信、导航、开健康码等任务往往要跨多个界面。现有语音助手多数停在对话或打开应用，"
      "不能稳定完成第三方界面上的点击与输入；脚本类工具又需要人工录制，老人用不了。")
    P(doc, "因此本系统要解决的核心问题是：在手机端把自然语言指令变成对真实界面的可控操作序列，"
      "并处理误操作风险、弱网条件和家属远程帮忙三类约束。")

    H(doc, "1.3 定义", 2)
    table(doc, ["名词", "含义"], [
        ["GUI Agent", "根据界面观测规划并执行 UI 动作的程序循环"],
        ["无障碍树", "Accessibility 返回的控件节点集合（文本、可点击性、坐标等）"],
        ["视觉模式", "无障碍树不可用时，用截图+坐标点击的回退方式"],
        ["Action-set", "预先编写的多步动作组，由状态机按阶段执行"],
        ["路由", "决定指令走模板、端侧 NLU、系统 Intent 还是 Agent"],
        ["KWS / VAD", "关键词唤醒 / 语音活动检测"],
        ["margin", "意图分类第一名与第二名概率差，用于抑制易混意图"],
    ])

    H(doc, "1.4 参考资料", 2)
    num(doc, [
        "第21届湖南省大学生计算机程序设计竞赛通知及材料要求",
        "Android Accessibility / Foreground Service / IME 文档",
        "ONNX Runtime、sherpa-onnx 相关说明",
        "本仓库源码、assets/nlu 配置与 tools/nlu 训练脚本",
    ])

    # ========== 2 ==========
    H(doc, "2 系统目标与设计约束", 1)
    H(doc, "2.1 功能目标", 2)
    num(doc, [
        "支持唤醒词、按键说话和文字输入三种下指令方式；",
        "能读取并操作本机第三方应用界面；",
        "常用指令尽量本地完成；复杂指令走大模型规划；",
        "发送、拨号前必须语音确认；",
        "支持家属协助码远程下发指令；",
        "覆盖通话、导航、设置、读通知、健康码/付款码、紧急求助等场景。",
    ])

    H(doc, "2.2 非功能目标", 2)
    table(doc, ["项目", "设计要求"], [
        ["正确性", "敏感动作有确认；不确定意图不盲目执行"],
        ["时延", "本地路径明显快于云端 Agent 路径"],
        ["弱网", "模板与端侧 NLU 在无大模型时仍可用"],
        ["可维护", "动作组、意图标签、工具表可扩展"],
        ["安全", "密钥本地存；上云文本可脱敏"],
    ])

    H(doc, "2.3 运行环境与约束", 2)
    table(doc, ["项目", "内容"], [
        ["平台", "Android 9.0+（minSdk 28 / targetSdk 36）"],
        ["语言", "Kotlin；协助服务端 Kotlin/JVM"],
        ["UI", "Jetpack Compose"],
        ["依赖权限", "麦克风、无障碍、悬浮窗等"],
        ["外部服务", "流式 ASR、大模型、高德、自建协助服务"],
        ["约束", "目标 App 若几乎不暴露控件树，需走视觉回退，成功率受界面影响"],
    ])

    # ========== 3 ==========
    H(doc, "3 总体设计", 1)
    H(doc, "3.1 设计思路", 2)
    P(doc, "系统采用“分层 + 级联路由 + 双执行引擎”的结构：")
    num(doc, [
        "分层：界面、运行时、决策、执行、外部服务分开，避免界面直接操作无障碍 API；",
        "级联路由：先便宜确定的路径，后昂贵不确定的路径；",
        "双执行引擎：系统 Intent / 动作组负责高频稳定任务，Agent 循环负责开放域任务；",
        "统一确认与守卫：不论哪条路径，敏感动作都进同一套确认逻辑。",
    ])

    H(doc, "3.2 逻辑结构", 2)
    pre(doc, [
        "用户(语音/文字) → 语音子系统(VAD/KWS/ASR/TTS)",
        "                 → AgentRuntime",
        "                 → CommandRouteResolver",
        "                       ├─ 模板匹配 ──────────────┐",
        "                       ├─ OfflineNluRouter ───────┤→ SystemIntentExecutor / 本地动作",
        "                       ├─ 系统意图解析 ───────────┤",
        "                       ├─ 预设指令 ───────────────┘",
        "                       └─ AgentOrchestrator",
        "                             ├─ 观测：UI树 / 截图",
        "                             ├─ 规划：LLM 工具调用 或 run_action_set",
        "                             ├─ 守卫：AgentActionGuard",
        "                             └─ 执行：AccessibilityGateway",
        "协助端 ──WebSocket── assist-server ──→ AgentRuntime（复用同上）",
    ])

    H(doc, "3.3 工程模块", 2)
    table(doc, ["模块", "职责"], [
        [":app", "主应用全部业务代码与资源"],
        [":assist-protocol", "协助消息结构"],
        [":assist-server", "配对、JWT、WebSocket 中继"],
        [":core:common", "公共代码"],
        [":uitreetest", "界面树调试"],
        ["tools/nlu", "意图训练、导出、hold-out 评测"],
        ["tools/ort-custom", "ONNX/ORT 相关优化产物"],
    ])

    H(doc, "3.4 包级模块一览", 2)
    table(doc, ["包", "主要类/对象", "职责"], [
        ["agent", "AgentRuntime, Orchestrator, RouteResolver, Guard, ToolRegistry, ActionSet", "决策与编排"],
        ["accessibility", "JoyAccessibilityService, Gateway", "读树与执行"],
        ["speech", "DoubaoAsrClient, JoyTtsSpeaker", "识别与播报"],
        ["wakeword", "WakeWordService, Sherpa检测器, SileroVadGate", "唤醒"],
        ["offline.nlu", "OfflineNluRouter, OnnxIntentClassifier", "端侧意图"],
        ["overlay / ime", "FloatingOverlayService, JoyInputMethodService", "悬浮与输入"],
        ["collaboration", "AssistSessionManager", "远程协助客户端"],
        ["caregiver / preset", "CaregiverSupportStore, PresetCommandStore", "照护与预设"],
        ["privacy", "PageContextRedactor", "脱敏"],
        ["system", "SystemIntentExecutor 等", "系统能力"],
        ["ui.cortana", "MainPivotScreen 等", "界面"],
    ])

    # ========== 4 创新 ==========
    H(doc, "4 关键技术创新点（设计层面）", 1)
    P(doc, "本章只写与实现直接对应的设计创新，便于评审查看“新在哪里、怎么做的”。")

    H(doc, "4.1 面向手机界面的 GUI Agent 执行闭环", 2)
    P(doc, "与“只聊天”的助手不同，本系统把大模型输出约束为工具白名单中的结构化动作"
      "（click/tap/type/send/scroll/open_app/finish/run_action_set 等），每一步都要基于"
      "当前界面观测，执行后再观察。设计上把“生成文本”改成“生成可执行 UI 操作”。")

    H(doc, "4.2 五级级联路由（成本与风险可控）", 2)
    P(doc, "CommandRouteResolver 固定顺序：模板 → 端侧 NLU → 系统意图 → 预设 → Agent。"
      "这样把“打开 WiFi”“几点了”等高频低歧义请求挡在本地，只有开放域跨界面任务才消耗"
      "ASR 之后的大模型规划，从设计上降低时延、费用和失败面。")

    H(doc, "4.3 端侧 Transformer 意图分类 + 三重门控", 2)
    P(doc, "端侧使用 chinese_roberta_L-4_H-256 微调后的 INT8 ONNX 分类器，最大长度 64。"
      "决策不是只看最高分，而是同时使用：")
    table(doc, ["门控参数", "配置值", "作用"], [
        ["auto_execute_threshold", "0.88", "高于此值才允许自动执行"],
        ["clarify_threshold", "0.75", "介于澄清与自动执行之间可追问"],
        ["margin_threshold", "0.28", "Top1-Top2 间隔过小则拒绝，防止近义意图误触发"],
    ])
    P(doc, "意图标签覆盖设置类、应用打开、闹钟日历、天气时间、拨号、导航回家、紧急求助、"
      "叫家人帮忙、健康码、付款码及 none 等。训练与 hold-out 评测脚本独立，避免只在训练集上自测。")

    H(doc, "4.4 本地唤醒链路：VAD → KWS → 预缓冲 ASR", 2)
    P(doc, "唤醒不是“检测到关键词就直接把空音频送给云端”，而是：Silero VAD 过滤非人声 →"
      "Sherpa Zipformer KWS 检出唤醒词 → 提示音 → 将唤醒前后缓冲音频接到流式 ASR。"
      "该设计减少误唤醒，并降低老人“喊完还要再说一遍”的成本。")

    H(doc, "4.5 敏感操作语音二次确认 + 动作守卫", 2)
    P(doc, "在工具协议中显式使用 needs_binary_confirm 标记。发送、拨号等动作由 AgentActionGuard"
      "改写为“先询问、再听短应答、再执行”。同时记录近期动作，抑制重复无效点击造成的死循环。"
      "确认逻辑对 Action-set 路径与自由 Agent 路径共用。")

    H(doc, "4.6 双通道界面感知", 2)
    P(doc, "主通道：无障碍树 + 文本打分点击（click）。回退通道：树为空或不可用时进入视觉模式，"
      "使用归一化坐标 tap（0~1000）。工具说明中对两种模式互斥约束（树可用禁止瞎估坐标；"
      "视觉模式强制 tap），减少模型乱选工具。")

    H(doc, "4.7 Action-set：固定多步状态机 + 窄域 LLM", 2)
    P(doc, "对微信发消息、淘宝搜索、地图导航等高频任务，不把每一步都交给主规划模型。"
      "设计 ActionSetDefinition（阶段图）与 ActionScript DSL，由 ActionSetFlowEngine 按阶段"
      "drain 动作。LLM 只需输出 run_action_set 及参数；阶段内可穿插 CapturePageTexts / AskLlm"
      "做局部决策。这样减少反复塞整棵 UI 树，提高路径稳定性。")
    table(doc, ["动作组示例", "参数", "用途"], [
        ["微信/IM 发送", "联系人、消息内容", "打开会话、输入、发送前确认"],
        ["淘宝搜索 / 搜索并打开", "关键词", "电商检索稳定路径"],
        ["地图导航", "目的地", "导航唤起"],
    ])

    H(doc, "4.8 家庭远程协助复用执行栈", 2)
    P(doc, "协助子系统只解决“连接与转发”，不复制一套点击引擎。家属指令进入老人端 AgentRuntime"
      "后与本地指令同路径执行，因此确认、守卫、动作组在远程场景仍然生效。协议与服务端独立成模块，便于演示部署。")

    H(doc, "4.9 隐私脱敏与输入稳定性工程", 2)
    P(doc, "上云前对页面上下文做 PageContextRedactor 脱敏；自动输入优先走 Joy IME，失败再降级。"
      "这些不是算法创新，但是银发场景落地必需的系统设计点，直接影响可用率与安全。")

    # ========== 5 详细流程 ==========
    H(doc, "5 关键处理流程设计", 1)

    H(doc, "5.1 语音主流程", 2)
    pre(doc, [
        "1  WakeWordService 采集音频",
        "2  SileroVadGate：非人声则丢弃",
        "3  Sherpa KWS：未命中则继续听；命中则发唤醒事件",
        "4  播放提示；WakeChainedAudioBridge 附带预缓冲",
        "5  DoubaoAsrClient 流式识别 → 文本",
        "6  文本进入 CommandRouteResolver（见 5.2）",
        "7  执行结果 → JoyTtsSpeaker 播报（可 barge-in）",
    ])

    H(doc, "5.2 路由判定流程", 2)
    pre(doc, [
        "输入: utterance",
        "if ElderTaskTemplateMatcher.hit: 执行模板绑定动作; return",
        "pred = OfflineNluRouter.predict(utterance)",
        "if pred 满足 auto_execute 且 margin 足够: 映射系统/本地动作; return",
        "if pred 落在澄清带: 追问关键槽位; return",
        "if 系统意图解析成功: SystemIntentExecutor; return",
        "if 预设指令命中: 执行预设; return",
        "else: AgentOrchestrator.run(utterance)",
    ])

    H(doc, "5.3 Agent 循环", 2)
    pre(doc, [
        "loop (step < maxSteps):",
        "  obs = gateway.observe()          # 树优先，必要时 screenshot",
        "  if 需要视觉模式: visionMode=true",
        "  action = llm.plan(goal, obs, toolsPrompt)",
        "  if action == run_action_set:",
        "      由 ActionSetFlowEngine 分阶段执行; 可提前 finish",
        "  action = AgentActionGuard.process(action)",
        "  if 需要确认: 语音询问; 等待 ASR; 否定则 break",
        "  result = gateway.execute(action)",
        "  记录 AgentStepRecord",
        "  if action.finished: break",
        "播报最终 message",
    ])

    H(doc, "5.4 确认子流程", 2)
    P(doc, "当动作带 needs_binary_confirm=true，运行时进入 waitingForUserConfirm。"
      "播报确认语后开启短听写；VoiceConfirmPhraseMatcher 判定肯定/否定；"
      "肯定则继续原动作，否定则取消并反馈。该状态与普通“等待用户补充信息”区分开。")

    H(doc, "5.5 远程协助流程", 2)
    pre(doc, [
        "老人端: 向 assist-server 申请/展示协助码",
        "家属端: 提交协助码建立会话（JWT）",
        "server: WebSocket 房间转发屏幕相关数据与控制消息",
        "老人端 AssistSessionManager 收指令 → AgentRuntime.submit",
        "执行路径与本地指令相同（含确认）",
    ])

    # ========== 6 模块详细 ==========
    H(doc, "6 模块详细设计", 1)

    H(doc, "6.1 AgentRuntime（运行时门面）", 2)
    P(doc, "职责：汇总权限检查、语音会话、唤醒事件、UI 状态（AgentUiState）、确认态、"
      "协助会话入口，对外提供 submitText / 语音结果回调等。不直接解析大模型 JSON，而是调度"
      "RouteResolver 与 Orchestrator。")
    table(doc, ["主要状态字段（逻辑）", "含义"], [
        ["conversation cards", "界面展示的对话/预览/确认卡片"],
        ["task steps", "当前任务步骤列表"],
        ["waitingForUserConfirm", "是否在敏感确认中"],
        ["confirmPrompt", "确认问题文本"],
        ["modelName 等", "当前模型与配置摘要"],
    ])

    H(doc, "6.2 CommandRouteResolver", 2)
    P(doc, "输入为字符串指令，输出为“已处理结果”或“转交 Agent”。设计要点是顺序固定、"
      "短路返回，避免同一句指令被多个执行器重复执行。模板匹配优先于模型，是为了保证"
      "“紧急呼救”“导航回家”等语句行为确定。")

    H(doc, "6.3 OfflineNluRouter / OnnxIntentClassifier", 2)
    P(doc, "处理流程：分词/WordPiece（与导出配置一致）→ ONNX 推理 → softmax 得各类概率 →"
      "取 Top1/Top2 算 margin → 与阈值比较 → IntentActionMapper 生成动作。"
      "模型资源：intent_classifier.onnx、vocab.txt、intent_labels.json、encoder_config.json。")
    P(doc, "当 model_type 被配置为降级方案时，可回退到更轻量分类器（工程预留），以适配低端机。")

    H(doc, "6.4 AgentOrchestrator 与工具层", 2)
    P(doc, "Orchestrator 维护一步或多步计划。AgentToolRegistry 定义工具名与提示词说明，"
      "并把动作分发给 AccessibilityActionDispatcher 或 SystemIntentExecutor。")
    P(doc, "工具分组：")
    num(doc, [
        "UI 原子操作：click, tap, type, send, scroll_*, swipe_down, back, home, wait；",
        "观测辅助：find_on_page, read_tree, query_page/diff/tree；",
        "应用与系统：list_apps, open_app, 各类 open_*_settings, dial_contact, set_alarm 等；",
        "银发能力：navigate_*, read_unread_messages, open_health_code/payment_code, emergency_help, ask_family_for_help；",
        "控制：finish（含 waiting_for_user / needs_binary_confirm）；",
        "复合：run_action_set。",
    ])
    P(doc, "规划返回 JSON 字段至少包括：action, target_text, input_text, message, finished,"
      " waiting_for_user, needs_binary_confirm。")

    H(doc, "6.5 Action-set 子系统", 2)
    P(doc, "核心类型：")
    table(doc, ["类型", "作用"], [
        ["ActionSetDefinition", "id、参数规格、uiLabel、phases 图"],
        ["ActionSetParams", "字符串参数表"],
        ["ActionSetFlowEngine", "根据已执行记录计算下一步 drain"],
        ["ActionSetDrain", "RunActions / CapturePageTexts / AskLlm / Done"],
        ["ActionScript DSL", "用线性脚本编译为 phase 图，降低编写成本"],
    ])
    P(doc, "设计原则：动作组执行期间尽量少调用主规划、少重复上传整树；缺参时宁可 AskLlm"
      "重试或中止，也不空目标硬点（AskPolicy）。")

    H(doc, "6.6 无障碍执行子系统", 2)
    P(doc, "JoyAccessibilityService 维护当前活动窗口节点。点击策略以目标文本与节点文本/"
      "内容描述的匹配打分为主，再校验可点击与可见性。输入路径：焦点控件 → Joy IME 注入 →"
      "失败则备选方案。微信等场景要求同时开启白名单读树组件。Gateway 对上层只暴露 observe/execute。")

    H(doc, "6.7 语音子系统", 2)
    table(doc, ["部件", "技术点"], [
        ["VAD", "Silero ONNX，门控是否送入 KWS"],
        ["KWS", "Sherpa-ONNX Zipformer，拼音单元，可配唤醒词与阈值"],
        ["ASR", "豆包流式；支持唤醒链式预缓冲"],
        ["TTS", "系统引擎；记录 prompt 并支持打断预滚动"],
        ["确认听写", "短窗口 ASR + 短语匹配"],
    ])

    H(doc, "6.8 安全与隐私模块", 2)
    P(doc, "AgentActionGuard：动作改写、确认注入、重复检测。PageContextRedactor：过滤验证码等"
      "敏感片段后再进入 LLM 上下文。密钥与偏好存 SharedPreferences，构建时 local.properties"
      "注入 BuildConfig，均不进入公开提交的源码内容。")

    H(doc, "6.9 协助子系统", 2)
    P(doc, "协议模块定义控制消息；服务端负责任房、鉴权、转发；客户端管理连接生命周期。"
      "角色分老人端/家属端。该子系统的创新不在网络协议本身，而在与 GUI Agent 执行栈的打通。")

    H(doc, "6.10 界面与辅助模块", 2)
    P(doc, "主界面四页：助手/设置/协作/关于。悬浮窗提供跨应用入口。IME 提升自动输入成功率。"
      "照护配置提供亲属别名、紧急文案、家地址，供模板与系统能力引用。")

    # ========== 7 接口 ==========
    H(doc, "7 接口设计", 1)
    H(doc, "7.1 外部接口", 2)
    table(doc, ["接口", "协议", "请求/要点", "响应/要点"], [
        ["大模型规划", "HTTPS JSON", "系统提示+工具说明+观测+用户目标", "单步动作 JSON"],
        ["豆包 ASR", "WebSocket", "音频帧流", "中间/最终文本"],
        ["高德", "HTTPS + Intent", "POI 查询；androidamap 深链", "候选点/拉起导航"],
        ["协助服务", "HTTP+WS", "协助码、JWT、指令帧", "会话事件、转发数据"],
        ["系统能力", "Intent", "拨号、设置页、闹钟等", "系统界面跳转"],
    ])

    H(doc, "7.2 内部关键接口（逻辑）", 2)
    table(doc, ["调用方", "被调方", "接口含义"], [
        ["UI", "AgentRuntime", "提交文本/开始听写/取消任务"],
        ["Runtime", "RouteResolver", "路由一条 utterance"],
        ["Orchestrator", "LLM Client", "plan(goal, obs)"],
        ["Orchestrator", "Gateway", "observe/execute"],
        ["Guard", "Runtime 确认通道", "请求二元确认结果"],
        ["AssistClient", "Runtime", "注入远程指令"],
    ])

    H(doc, "7.3 用户接口", 2)
    num(doc, [
        "语音：唤醒、按钮听写、确认回答、打断播报；",
        "文字：输入框；",
        "快捷：建议指令芯片；",
        "状态：步骤列表、对话卡片、确认卡片；",
        "协同：协助码显示与输入。",
    ])

    # ========== 8 数据 ==========
    H(doc, "8 数据结构设计", 1)
    H(doc, "8.1 核心内存结构", 2)
    table(doc, ["结构", "主要字段", "用途"], [
        ["AgentAction", "action, targetText, inputText, message, finished, 确认标志", "一步操作"],
        ["AgentStepRecord", "动作、结果、时间", "循环控制与动作组对齐"],
        ["界面观测", "节点列表/文本摘要/截图", "规划输入"],
        ["ActionSetParams", "Map<String,String>", "动作组参数"],
        ["IntentPrediction", "label, probs, margin", "端侧路由"],
        ["AssistMessage", "类型、载荷", "远程协同"],
    ])

    H(doc, "8.2 持久化设计", 2)
    P(doc, "客户端不以业务库为中心，采用 SharedPreferences 分类存储：API、唤醒、会话、照护、"
      "预设、协助绑定、主题等。assets 存放模型与词表。assist-server 用 H2 存房间/会话元数据。"
      "若竞赛材料需要《数据库设计说明书》，可声明“非必要”并以本章代替逻辑数据说明。")

    H(doc, "8.3 主要数据流", 2)
    num(doc, [
        "音频流：Mic → VAD → KWS → ASR；",
        "文本流：ASR/UI/远程 → Router → 执行器；",
        "观测流：Accessibility → Orchestrator/LLM；",
        "动作流：LLM/动作组 → Guard → Gateway → App；",
        "协同流：家属 → Server → 老人 Runtime。",
    ])

    # ========== 9 运行 ==========
    H(doc, "9 运行设计", 1)
    H(doc, "9.1 运行组合", 2)
    table(doc, ["模式", "需要启用的部件"], [
        ["最小演示", "App + 无障碍 + 麦克风 + ASR/LLM 密钥"],
        ["完整体验", "最小 + 唤醒FGS + 悬浮窗 + 微信组件 + IME + 照护配置"],
        ["远程协助", "完整体验 + assist-server 网络可达"],
        ["弱网本地", "模板/端侧 NLU/系统 Intent（不依赖 LLM）"],
    ])

    H(doc, "9.2 控制策略", 2)
    num(doc, [
        "前台服务保活唤醒与悬浮能力，并显示通知；",
        "Agent 最大步数限制；",
        "确认态互斥：确认完成前不继续敏感动作；",
        "协助断连即停转发；",
        "协程取消用于用户主动停止任务。",
    ])

    H(doc, "9.3 性能相关设计", 2)
    P(doc, "VAD 减少无效 KWS；INT8 意图模型控制包体与算力；路由短路减少 LLM 调用；"
      "Action-set 减少多轮规划；观测上传前做摘要/脱敏，控制上下文体积。")

    # ========== 10 出错 ==========
    H(doc, "10 出错处理设计", 1)
    table(doc, ["故障", "检测", "处理"], [
        ["无障碍关闭", "服务未绑定", "提示开启并中止代操作"],
        ["麦克风拒绝", "权限检查", "关闭语音入口并提示"],
        ["ASR/LLM 失败", "网络/HTTP 错误", "提示重试；可回退本地能力"],
        ["节点找不到", "匹配分过低", "滚动/换工具/视觉回退；超限失败说明"],
        ["意图混淆", "margin 不足", "不自动执行，转澄清或 Agent"],
        ["用户拒绝确认", "否定短语", "取消动作并播报"],
        ["动作组缺参", "AskPolicy", "窄域追问或中止，禁止空点"],
        ["协助失败", "连接/鉴权错误", "提示检查码与服务"],
    ])

    # ========== 11 安全 ==========
    H(doc, "11 安全保密设计", 1)
    num(doc, [
        "敏感动作确认；",
        "动作循环抑制；",
        "上下文脱敏；",
        "密钥本地化与 .gitignore；",
        "高权限能力明示用途；",
        "远程协助建议可信网络，生产需再加强传输与审计；",
        "紧急求助不能替代 120 等公共服务。",
    ])

    # ========== 12 界面 ==========
    H(doc, "12 界面结构设计摘要", 1)
    table(doc, ["页面", "元素", "对应模块"], [
        ["助手", "对话、步骤、建议芯片、输入/语音", "Runtime / UI"],
        ["设置", "唤醒、API、开关", "各 ConfigStore"],
        ["协作", "角色、协助码、远程指令", "collaboration"],
        ["关于", "版本、模型、记忆摘要", "UI"],
        ["系统层", "悬浮窗、通知", "overlay / FGS"],
    ])
    P(doc, "视觉规范与截图见《软件界面设计书》。")

    # ========== 13 维护测试 ==========
    H(doc, "13 维护与测试设计", 1)
    H(doc, "13.1 维护", 2)
    P(doc, "第三方 App 改版时优先改 Action-set 与点击匹配；意图变化时重训 nlu 并跑 hold-out；"
      "模型供应商变更只改配置与 Client 适配。uitreetest 与树日志用于定位点击失败。")

    H(doc, "13.2 测试用例（设计阶段列出）", 2)
    table(doc, ["编号", "场景", "期望"], [
        ["T1", "打开 WiFi/放大字体", "走本地路径成功"],
        ["T2", "导航回家", "调起地图"],
        ["T3", "打电话确认-同意", "确认后拨号"],
        ["T4", "打电话确认-拒绝", "不拨号并提示取消"],
        ["T5", "微信发送动作组", "进入会话并在发送前确认"],
        ["T6", "无障碍关闭", "有明确提示"],
        ["T7", "断网时本地意图", "仍可完成部分设置类指令"],
        ["T8", "协助码远程指令", "老人端执行且确认仍生效"],
        ["T9", "意图近义句", "margin 不足时不误自动执行"],
        ["T10", "树空界面", "进入视觉 tap 路径或可失败降级"],
    ])

    # ========== 14 ==========
    H(doc, "14 小结", 1)
    P(doc, "JoyForOld 的系统设计围绕“语音到界面动作”的完整链路展开：用五级路由控制成本与风险，"
      "用端侧 NLU 三重门控处理常用意图，用 Agent 工具循环处理开放域任务，用 Action-set 提高"
      "高频路径稳定性，用确认守卫与远程协助补齐安全与家庭托底。本文给出了与源码一致的模块、"
      "流程、接口与数据设计，可作为竞赛评审与后续详细设计的基础。")
    P(doc, "（正文结束）", align="center", indent=False, before=18)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for i in range(1, 4):
        doc.styles[f"Heading {i}"].font.color.rgb = RGBColor(0, 0, 0)

    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, side, Cm(2.5))

    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("The 21th Hunan Collegiate Programming Contest")
    font(hr, size=9, color=RGBColor(0x66, 0x66, 0x66))
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("注意：文档内容要自动目录，双面打印，白色纸质胶装，一式2份；打印时请删除本页脚。")
    font(fr, size=8, cn="宋体", color=RGBColor(0x99, 0x99, 0x99))

    cover(doc)
    toc(doc)
    body(doc)

    saved = []
    for path in OUTS:
        try:
            doc.save(str(path))
            saved.append(path)
        except PermissionError:
            print("Locked:", path.name)
    if not saved:
        alt = DOCS / "JoyForOld-System-Design-Specification-full.docx"
        doc.save(str(alt))
        saved.append(alt)
    for p in saved:
        print("Saved", p, p.stat().st_size)


if __name__ == "__main__":
    main()
