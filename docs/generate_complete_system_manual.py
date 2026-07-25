# -*- coding: utf-8 -*-
"""
JoyForOld 系统设计说明书（合订本）
文风参考：软工示例文档的章节体例 + 竞赛可读表述
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOCS = Path(__file__).resolve().parent
OUTS = [
    DOCS / "JoyForOld-System-Manual-Complete.docx",
]


def font(run, size=12, bold=False, cn="宋体", en="Times New Roman", color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = en
    if color:
        run.font.color.rgb = color
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), cn)


def P(doc, text, size=12, bold=False, align="justify", indent=True, before=0, after=6, cn="宋体"):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.5
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf.first_line_indent = Cm(0.74) if indent and align == "justify" else Cm(0)
    r = p.add_run(text)
    font(r, size=size, bold=bold, cn=cn)
    return p


def H(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for r in p.runs:
        font(r, size={1: 16, 2: 14, 3: 12, 4: 12}[min(level, 4)], bold=True, cn="黑体")
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)


def num(doc, items):
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"（{i}）{t}")
        font(r, size=12)


def bullets(doc, items):
    for t in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(t)
        font(r, size=12)


def cell(c, text, bold=False, center=False, size=10.5):
    c.text = ""
    p = c.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    font(r, size=size, bold=bold)


def table(doc, headers, rows, size=10.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell(t.rows[0].cells[i], h, bold=True, center=True, size=size)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            cell(t.rows[ri + 1].cells[ci], v, size=size)
    doc.add_paragraph()


def pre(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(line)
        font(r, size=9.5, cn="宋体", en="Courier New")


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    P(doc, "第21届湖南省大学生计算机程序设计竞赛", size=16, bold=True, align="center", indent=False, cn="黑体")
    P(doc, "——应用开发类竞赛（2025）", size=14, bold=True, align="center", indent=False, cn="黑体", after=20)
    for _ in range(2):
        doc.add_paragraph()
    P(doc, "系  统  设  计  说  明  书", size=28, bold=True, align="center", indent=False, cn="黑体", after=8)
    P(doc, "（合订本）", size=14, align="center", indent=False, cn="黑体", after=20)
    P(doc, "含：功能需求 · 概要设计 · 详细设计 · 数据存储 · 界面设计 · 使用与安装说明", size=11, align="center", indent=False, after=24)
    for k, v in [
        ("作品名称：", "JoyForOld——面向老年用户的手机语音操控助手"),
        ("作品类别：", "移动终端开发类"),
        ("作　　者：", "（请填写参赛队员姓名）"),
        ("指导老师：", "（请填写指导教师姓名）"),
        ("单　　位：", "（请填写学校名称）"),
        ("文档版本：", "V1.0"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(k)
        font(r1, size=13, bold=True, cn="黑体")
        r2 = p.add_run(v)
        font(r2, size=13)
    for _ in range(3):
        doc.add_paragraph()
    P(doc, "二零二五年   月   日", size=14, align="center", indent=False)
    page_break(doc)


def toc_page(doc):
    P(doc, "目    录", size=16, bold=True, align="center", indent=False, cn="黑体", after=12)
    P(doc,
      "请在 Microsoft Word 中于下方目录域上右键选择“更新域”，以生成带页码的自动目录。"
      "胶装打印时请使用官方封面模板，双面打印，白色纸质胶装，一式两份，并删除页脚提示文字。",
      size=10.5, indent=False, after=12)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run._r.append(run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"}))
    run2 = paragraph.add_run()
    instr = run2._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = r' TOC \o "1-3" \h \z \u '
    run2._r.append(instr)
    run3 = paragraph.add_run()
    run3._r.append(run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"}))
    run4 = paragraph.add_run("（目录域：请更新）")
    font(run4, size=12, color=RGBColor(0x66, 0x66, 0x66))
    run5 = paragraph.add_run()
    run5._r.append(run5._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"}))
    page_break(doc)


def part_meta(doc, title):
    """各分册开头的统一“文档介绍”短段。"""
    H(doc, title, 1)


def body(doc):
    # ========== 总述 ==========
    H(doc, "前言", 1)
    P(doc,
      "本合订本为参赛作品 JoyForOld 的系统设计说明书，依据湖南省大学生计算机程序设计竞赛"
      "应用开发类材料要求整理而成。全书按软件工程常见体例编排：先明确需求，再给出总体结构，"
      "继而说明关键模块的处理过程，并附以数据、界面与使用安装说明，以便评审专家全面了解"
      "本系统“做什么、如何组织、如何实现、如何使用”。")
    P(doc,
      "文中所述功能与结构均以当前工程实现为依据。作者、指导教师与单位信息请于封面处据实填写。"
      "若正文表述与源码细节偶有出入，以源码为准，并应在修订版中同步更正。")

    # ========== 第一部分 需求 ==========
    H(doc, "第一部分  功能需求说明书", 1)

    H(doc, "1.1 文档介绍", 2)
    H(doc, "1.1.1 编写目的", 3)
    P(doc, "本部分用于界定 JoyForOld 的功能范围与质量要求，作为开发、测试与验收的共同依据，"
      "亦便于评审专家在短时间内把握系统能力边界。")
    H(doc, "1.1.2 读者对象", 3)
    P(doc, "指导教师、开发与测试人员、竞赛评审专家，以及需要了解本系统能力的相关人员。")
    H(doc, "1.1.3 术语与缩写", 3)
    table(doc, ["术语", "说明"], [
        ["无障碍服务", "Android Accessibility Service，用于读取界面并模拟操作"],
        ["唤醒词", "用于启动听写的固定口语短语"],
        ["ASR", "自动语音识别"],
        ["TTS", "语音合成播报"],
        ["NLU", "本系统中指端侧意图分类与简单槽位解析"],
        ["Agent", "根据界面状态规划并执行操作的程序循环"],
        ["Action-set", "预先编排的多步动作组"],
        ["协助码", "老人端与家属端建立远程协助所用的短码"],
    ])
    H(doc, "1.1.4 参考资料", 3)
    num(doc, [
        "第21届湖南省大学生计算机程序设计竞赛通知及提交要求",
        "Android 开发者文档（Accessibility、前台服务、输入法等）",
        "本项目源码、README 与 assets 中的模型配置",
    ])

    H(doc, "1.2 项目介绍", 2)
    H(doc, "1.2.1 项目说明", 3)
    table(doc, ["项目", "内容"], [
        ["项目名称", "JoyForOld"],
        ["作品定位", "面向老年用户的 Android 语音与界面操控助手"],
        ["运行平台", "Android 9.0 及以上智能手机"],
        ["主要用户", "老年用户；家属/照护者（远程协助）"],
    ])
    H(doc, "1.2.2 项目背景", 3)
    P(doc,
      "智能手机已成为老年人联系亲友、出行就医与办理日常事务的重要工具。然而，许多应用界面"
      "层级较深、操作步骤较多，老年人独立完成“打电话、发消息、导航、打开健康码”等任务仍有困难。"
      "常见语音助手擅长问答与部分系统控制，对第三方应用内的连续点击与输入支持有限；子女通过"
      "视频口头指导，又往往因界面不同步而效率不高。")
    P(doc,
      "据此，本项目希望提供一种更贴近生活任务的辅助方式：用户用自然语言表达意图，系统在授权"
      "前提下读取当前界面并代为操作；对发送、拨号等重要步骤先行确认；必要时由家属远程协助。")

    H(doc, "1.2.3 项目目标", 3)
    num(doc, [
        "支持语音与文字下达生活类指令，并完成可观察的操作结果；",
        "对高频、表述清晰的指令尽量在本地快速处理；",
        "对复杂、跨界面任务，结合大模型与界面观测分步执行；",
        "降低误发、误拨风险；",
        "支持家属在需要时远程帮忙。",
    ])

    H(doc, "1.2.4 用户角色", 3)
    table(doc, ["角色", "特征", "主要期望"], [
        ["老年用户", "偏好口语，担心点错", "少记步骤、能办成事、关键操作有确认"],
        ["家属/照护者", "可能不在身边", "能连接、能下发指令、可配置亲情信息"],
        ["配置人员", "首次安装与演示", "权限与密钥配置清楚、场景可复现"],
    ])

    H(doc, "1.3 需求编号规则与模块划分", 2)
    P(doc, "功能需求编号格式为：JFO-F-模块缩写-序号。例如 JFO-F-VOC-001 表示语音相关第 1 条。")
    table(doc, ["模块缩写", "名称", "说明"], [
        ["VOC", "语音交互", "唤醒、识别、播报、确认听写"],
        ["RTE", "指令处理", "路由、本地意图、Agent 规划"],
        ["ACT", "界面执行", "无障碍点击、输入、打开应用等"],
        ["SAF", "安全确认", "敏感操作确认与异常抑制"],
        ["COL", "远程协助", "协助码与远端指令"],
        ["CFG", "配置与照护", "密钥、联系人、预设、家地址等"],
        ["UI", "人机界面", "主界面、悬浮窗、建议指令"],
    ])

    H(doc, "1.4 功能性需求", 2)

    H(doc, "1.4.1 语音交互（VOC）", 3)
    table(doc, ["编号", "需求名称", "需求描述"], [
        ["JFO-F-VOC-001", "本地唤醒", "在授权后可监听唤醒词；命中后进入听写，并给出可感知提示"],
        ["JFO-F-VOC-002", "语音听写", "将用户语音转为文本指令；支持唤醒后的连续表达"],
        ["JFO-F-VOC-003", "结果播报", "以语音播报执行结果、询问与提示；允许用户打断"],
        ["JFO-F-VOC-004", "确认听写", "在敏感确认阶段，短时听取肯定或否定回答"],
        ["JFO-F-VOC-005", "文字入口", "允许用户以文本框提交指令，作为语音的补充"],
    ])

    H(doc, "1.4.2 指令处理（RTE）", 3)
    table(doc, ["编号", "需求名称", "需求描述"], [
        ["JFO-F-RTE-001", "模板匹配", "对“打电话给××”“导航回家”“紧急呼救”等高频语句给出确定行为"],
        ["JFO-F-RTE-002", "端侧意图", "在设备本地识别常用意图（设置、天气、拨号等），高把握时直接执行"],
        ["JFO-F-RTE-003", "系统能力", "调用拨号、设置页、闹钟、日历、地图深链等系统能力"],
        ["JFO-F-RTE-004", "预设指令", "支持家属配置的短语别名"],
        ["JFO-F-RTE-005", "开放域规划", "对未能本地处理的指令，结合界面信息分步规划并执行"],
        ["JFO-F-RTE-006", "动作组", "对微信发送、地图导航、电商搜索等提供相对稳定的多步路径"],
    ])

    H(doc, "1.4.3 界面执行（ACT）", 3)
    table(doc, ["编号", "需求名称", "需求描述"], [
        ["JFO-F-ACT-001", "读界面", "读取当前窗口控件信息；必要时获取截图"],
        ["JFO-F-ACT-002", "点击与手势", "按文本或坐标点击，支持滑动、返回、回桌面等"],
        ["JFO-F-ACT-003", "自动输入", "向输入框写入文字；优先使用辅助输入法"],
        ["JFO-F-ACT-004", "打开应用", "按应用名启动本机应用"],
        ["JFO-F-ACT-005", "微信读树支持", "在需要时配合专用无障碍组件，提高特定应用读界面成功率"],
    ])

    H(doc, "1.4.4 安全确认（SAF）", 3)
    table(doc, ["编号", "需求名称", "需求描述"], [
        ["JFO-F-SAF-001", "发送确认", "发送消息前语音询问，未获同意不得发送"],
        ["JFO-F-SAF-002", "拨号确认", "拨出电话前语音询问，未获同意不得拨出"],
        ["JFO-F-SAF-003", "异常抑制", "对短时间重复无效操作予以限制，避免长时间空转"],
        ["JFO-F-SAF-004", "上云脱敏", "向云端提供界面摘要前，对明显敏感片段做处理"],
    ])

    H(doc, "1.4.5 远程协助（COL）", 3)
    table(doc, ["编号", "需求名称", "需求描述"], [
        ["JFO-F-COL-001", "协助码配对", "老人端展示协助码，家属端输入后建立会话"],
        ["JFO-F-COL-002", "远端指令", "家属下发的指令在老人手机执行，并尽量沿用本地安全确认"],
        ["JFO-F-COL-003", "会话结束", "任一方断开后停止转发"],
    ])

    H(doc, "1.4.6 配置与界面（CFG / UI）", 3)
    table(doc, ["编号", "需求名称", "需求描述"], [
        ["JFO-F-CFG-001", "接口配置", "可配置语音识别与大模型相关密钥"],
        ["JFO-F-CFG-002", "照护信息", "可配置亲属称呼与电话、紧急文案、家地址等"],
        ["JFO-F-CFG-003", "唤醒配置", "可设置唤醒词及相关参数"],
        ["JFO-F-UI-001", "主界面", "提供助手、设置、协作、关于等分区"],
        ["JFO-F-UI-002", "建议指令", "首页提供常用语句快捷入口"],
        ["JFO-F-UI-003", "悬浮入口", "在其他应用上层提供助手入口"],
        ["JFO-F-UI-004", "进度可见", "任务执行过程中展示步骤或状态，避免“黑盒感”"],
    ])

    H(doc, "1.5 非功能性需求", 2)
    table(doc, ["类别", "要求"], [
        ["正确性", "敏感操作必须确认；意图不明确时不贸然自动执行"],
        ["响应", "本地可处理指令应明显快于需多轮云端规划的任务"],
        ["弱网", "部分设置类、模板类指令在大模型不可用时仍可尝试完成"],
        ["兼容", "真机 Android 9.0 及以上；依赖系统无障碍能力"],
        ["可维护", "意图标签、动作组与工具表应可扩展"],
        ["安全", "密钥本地保存；高权限能力需用户知情开启"],
    ])

    H(doc, "1.6 约束与假设", 2)
    num(doc, [
        "用户须主动授予无障碍、麦克风等必要权限；",
        "目标应用若几乎不暴露控件信息，需依赖截图坐标等回退手段，成功率受界面影响；",
        "云端语音识别与大模型需网络与合法密钥；",
        "紧急求助依赖预配置联系人与通信条件，不能替代公共急救服务。",
    ])

    # ========== 第二部分 概要设计 ==========
    H(doc, "第二部分  概要设计说明书", 1)

    H(doc, "2.1 文档说明", 2)
    P(doc, "本部分说明系统的总体结构、模块职责、主要处理流程、接口关系与运行方式，"
      "为详细设计与编码提供框架。读者对象为开发人员与评审专家。")

    H(doc, "2.2 设计原则", 2)
    num(doc, [
        "先分流、后执行：能本地确定的指令，不默认进入大模型规划；",
        "观测与行动成对出现：规划所依据的信息来自当前界面，执行后再核对；",
        "安全优先于完全自动：发送、拨号等步骤须确认；",
        "稳定路径与灵活规划并用：高频流程可脚本化，开放任务走通用循环；",
        "远程协助复用本地执行能力，避免另造一套操作引擎。",
    ])

    H(doc, "2.3 系统结构", 2)
    P(doc, "系统在工程上分为主应用、协助协议、协助服务端及工具脚本等模块；在逻辑上分为"
      "界面层、运行时层、决策层、执行层与外部服务层。")
    pre(doc, [
        "┌──────────────────────────────────────┐",
        "│ 界面层：主界面 / 悬浮窗 / 通知 / 播报   │",
        "├──────────────────────────────────────┤",
        "│ 运行时层：AgentRuntime（会话与状态）    │",
        "├──────────────────────────────────────┤",
        "│ 决策层：路由 / 端侧NLU / Agent / 动作组 │",
        "├──────────────────────────────────────┤",
        "│ 执行层：无障碍网关 / 系统Intent / IME   │",
        "├──────────────────────────────────────┤",
        "│ 外部：ASR · 大模型 · 地图 · 协助服务    │",
        "└──────────────────────────────────────┘",
    ])

    H(doc, "2.4 工程模块一览", 2)
    table(doc, ["模块", "职责"], [
        [":app", "Android 主程序：界面、语音、决策、无障碍客户端等"],
        [":assist-protocol", "老人端与家属端消息约定"],
        [":assist-server", "协助中继：配对、鉴权、转发"],
        [":core:common", "公共代码"],
        [":uitreetest", "界面树调试辅助"],
        ["tools/nlu", "端侧意图训练、导出与评测"],
    ])

    H(doc, "2.5 逻辑模块划分", 2)
    table(doc, ["模块", "主要职责", "代表实现"], [
        ["语音唤醒", "VAD、唤醒词检出、前台监听", "WakeWordService 等"],
        ["语音识别与播报", "流式 ASR、TTS、打断", "DoubaoAsrClient、JoyTtsSpeaker"],
        ["指令路由", "按优先级选择处理路径", "CommandRouteResolver"],
        ["端侧意图", "本地分类与门控", "OfflineNluRouter"],
        ["智能体编排", "观察—规划—执行循环", "AgentOrchestrator"],
        ["动作组", "固定多步任务", "ActionSet 及相关 DSL"],
        ["无障碍执行", "读树、点击、输入、截图", "JoyAccessibilityService"],
        ["安全守卫", "确认注入、防空转", "AgentActionGuard"],
        ["远程协助", "配对与指令注入", "AssistSessionManager"],
        ["照护配置", "亲属、紧急文案、家地址", "CaregiverSupportStore 等"],
        ["人机界面", "分区导航与状态展示", "ui/cortana"],
    ])

    H(doc, "2.6 主处理流程", 2)
    P(doc, "用户通过唤醒、按键说话或文本输入形成指令文本后，系统按下列次序处理：")
    pre(doc, [
        "指令文本",
        "  → 模板匹配？是 → 执行绑定动作 → 播报 → 结束",
        "  → 端侧意图高置信？是 → 本地/系统执行 → 播报 → 结束",
        "  → 系统意图/预设？是 → 执行 → 播报 → 结束",
        "  → 否则进入 Agent：",
        "        观察界面（树优先，必要时截图）",
        "        → 大模型给出动作或 run_action_set",
        "        → 安全检查（敏感则确认）",
        "        → 无障碍/系统执行",
        "        → 未完成则继续，完成则播报结束",
    ])

    H(doc, "2.7 关键技术选型", 2)
    table(doc, ["方面", "选型", "简要理由"], [
        ["语言与界面", "Kotlin、Jetpack Compose", "适合现代 Android 应用结构"],
        ["依赖注入/异步", "Koin、Coroutines", "便于组织语音与网络并发"],
        ["界面操控", "Accessibility", "无需 Root 即可跨应用操作"],
        ["唤醒与 VAD", "sherpa-onnx、Silero", "可端侧部署"],
        ["端侧意图", "ONNX Runtime + 小规模中文模型 INT8", "体积与速度可接受"],
        ["云端 ASR", "豆包流式识别", "中文连续识别体验较好"],
        ["云端规划", "可配置大模型接口", "便于演示与替换"],
        ["协助服务", "Ktor、WebSocket、JWT、H2", "轻量、可自建"],
        ["地图", "高德查询与深链", "符合国内使用习惯"],
    ])

    H(doc, "2.8 接口概要", 2)
    H(doc, "2.8.1 外部接口", 3)
    table(doc, ["接口", "方式", "用途"], [
        ["大模型", "HTTPS", "根据目标与界面给出下一步动作"],
        ["语音识别", "WebSocket", "语音转文字"],
        ["地图", "HTTPS / Intent", "地点解析与导航唤起"],
        ["协助服务", "HTTP + WebSocket", "配对与消息转发"],
        ["系统能力", "Intent 等", "拨号、设置、闹钟等"],
        ["无障碍框架", "系统回调与动作 API", "读界面、执行手势"],
    ])
    H(doc, "2.8.2 用户接口", 3)
    P(doc, "用户可通过唤醒词、语音按钮、文本框与建议语句下发指令；通过语音回答确认问题；"
      "在协作页进行协助码操作；通过悬浮窗在其他应用中唤起助手。")

    H(doc, "2.9 运行设计概要", 2)
    table(doc, ["运行模式", "需要具备的条件"], [
        ["基本演示", "主应用、无障碍、麦克风、云端密钥"],
        ["完整体验", "基本演示 + 唤醒服务 + 悬浮窗 + 微信支持组件 + 输入法 + 照护配置"],
        ["远程协助", "完整体验 + 可访问的协助服务端"],
        ["弱网本地", "模板与端侧意图、系统能力（不依赖大模型）"],
    ])

    H(doc, "2.10 设计要点（供评审把握创新与难点）", 2)
    P(doc, "下列要点体现本系统在设计上的着力处，详细算法见第三部分。")
    num(doc, [
        "将自然语言理解的结果落实为界面动作，而不仅停留在对话回复；",
        "以多级路由控制云端调用范围，兼顾速度与费用；",
        "端侧意图采用置信度与类别间隔双重约束，减少误触发；",
        "无障碍树与截图坐标两种观测方式互相补充；",
        "高频任务以动作组固化步骤，降低对临时规划的依赖；",
        "敏感操作统一确认，远程指令亦复用同一执行与安全路径。",
    ])

    # ========== 第三部分 详细设计 ==========
    H(doc, "第三部分  详细设计说明书", 1)

    H(doc, "3.1 说明", 2)
    P(doc, "本部分对关键模块给出职责、输入输出、处理逻辑与协作关系。为控制篇幅，一般界面控件级"
      "描述从略，重点放在语音链路、路由、意图、Agent、动作组、无障碍与协助等核心路径。")

    H(doc, "3.2 运行时模块（AgentRuntime）", 2)
    P(doc, "职责：作为对外门面，承接界面与语音事件，维护会话状态，调度路由与编排，管理确认态"
      "与协助入口。不负责具体点击算法，也不直接解析模型原始报文（由客户端与编排器处理）。")
    table(doc, ["输入", "输出/副作用"], [
        ["用户文本、ASR 最终结果、唤醒事件、远程指令", "更新界面状态、触发执行、发起播报或确认"],
        ["取消/停止请求", "中止当前任务协程"],
    ])
    P(doc, "状态中与评审相关的字段包括：对话卡片、任务步骤、是否等待确认、确认问句、当前模型信息等。"
      "确认态与普通“等待用户补充信息”相区分，以免把开放问答误当成发送确认。")

    H(doc, "3.3 指令路由模块", 2)
    P(doc, "职责：按固定优先级选择处理路径，并短路返回，避免同一指令被多次执行。")
    P(doc, "处理顺序：")
    num(doc, [
        "长辈任务模板匹配；",
        "端侧意图分类（含门控）；",
        "本地系统意图解析；",
        "预设指令；",
        "以上皆非，则交 Agent 编排器。",
    ])
    P(doc, "设计说明：模板置于最前，是为了保证“紧急呼救”“导航回家”等语句行为稳定；"
      "端侧意图次之，用于消化大量设置与查询类短句；Agent 置于最后，用于消化开放表述。")

    H(doc, "3.4 端侧意图模块", 2)
    P(doc, "职责：在设备本地判断用户意图，并在把握足够时映射为可执行动作。")
    table(doc, ["项目", "内容"], [
        ["基座", "中文 RoBERTa-Mini（4 层 / 256 维）微调"],
        ["部署", "INT8 ONNX，最大长度 64"],
        ["资源", "intent_classifier.onnx、vocab.txt、intent_labels.json、encoder_config.json"],
        ["自动执行阈值", "0.88"],
        ["澄清阈值", "0.75"],
        ["间隔阈值 margin", "0.28（Top1 与 Top2 之差）"],
    ])
    P(doc, "处理步骤：文本规范化 → 分词/编码（与导出配置一致）→ 推理得各类概率 → 计算 margin →"
      "与阈值比较 → 通过则经规则补全槽位并映射动作，否则返回空交由后续路由。")
    P(doc, "主要意图类别包括：各类设置入口、打开应用、闹钟与日程、相机相册、天气与时间、"
      "拨打联系人、导航回家、紧急求助、呼叫家人、健康码、付款码，以及 none。")
    P(doc, "训练与评测脚本位于 tools/nlu，并配备与训练集隔离的 hold-out 样本，以减少“只在训练句上好看”的偏差。")

    H(doc, "3.5 语音唤醒与识别模块", 2)
    P(doc, "唤醒链路：麦克风采集 → Silero VAD 判断是否为人声 → Sherpa Zipformer 关键词检出 →"
      "提示反馈 → 将预缓冲音频接入流式识别。VAD 用于降低空闲误触发；预缓冲用于减少唤醒后的表达断裂。")
    P(doc, "识别采用云端流式 ASR；播报采用系统 TTS，并记录当前提示语以便打断处理。"
      "确认阶段使用较短的听写窗口，并结合肯定/否定短语匹配。")

    H(doc, "3.6 Agent 编排与工具层", 2)
    P(doc, "编排器维护步数上限，循环执行：获取观测 → 调用大模型（或推进动作组）→ 守卫处理 → 执行 → 记录步骤。")
    P(doc, "工具按用途分为：")
    num(doc, [
        "界面原子操作：click、tap、type、send、scroll、swipe、back、home、wait；",
        "观测辅助：find_on_page、read_tree、query_page 等；",
        "应用与系统：open_app、各类设置、dial_contact、set_alarm、navigate_* 等；",
        "银发相关：健康码/付款码、读未读消息、紧急求助、呼叫家人等；",
        "控制类：finish（可标记等待用户、是否二元确认）；",
        "复合类：run_action_set。",
    ])
    P(doc, "规划结果以 JSON 描述单步动作。当无障碍树可用时，优先 click 按文本点击；树不可用时进入视觉模式，"
      "使用归一化坐标 tap（0～1000）。二者在提示词中互相约束，以减少工具误用。")

    H(doc, "3.7 动作组（Action-set）", 2)
    P(doc, "对微信发消息、淘宝搜索、地图导航等任务，采用阶段式定义：参数表 + 阶段图。"
      "引擎根据已执行记录计算下一步，可能输出一批 UI 动作、采集页内文本、窄域询问模型，或结束。")
    P(doc, "设计意图：把“固定步骤”留在本地执行，减少反复把整棵界面树交给主规划模型；"
      "参数不全时允许有限次追问，避免在空目标上强行点击。")
    table(doc, ["动作组", "主要参数", "说明"], [
        ["IM/微信发送", "联系人、消息内容", "进入会话、输入，发送前进入确认"],
        ["淘宝搜索 / 搜索并打开", "关键词", "电商检索路径"],
        ["地图导航", "目的地", "导航唤起"],
    ])

    H(doc, "3.8 无障碍执行模块", 2)
    P(doc, "服务端读取当前窗口节点，按目标文本与节点文本、内容描述等进行匹配打分，结合可点击性"
      "与可见性选择目标。输入优先经 Joy 输入法写入；失败时再尝试其他方式。"
      "对微信等场景，需同时开启主服务与支持组件。上层通过网关调用，以隔离服务生命周期细节。")

    H(doc, "3.9 安全确认模块", 2)
    P(doc, "当动作被标记为需要二元确认，或守卫识别出发送/拨号等风险时，系统改写执行流："
      "先播报问句并进入确认态，再根据听写结果决定继续或取消。同时对短期内重复失败的同类动作"
      "进行抑制，以免界面卡住时无意义重试。")

    H(doc, "3.10 远程协助模块", 2)
    P(doc, "老人端申请并展示协助码；家属端提交协助码建立会话。服务端完成房间管理、令牌校验与"
      "WebSocket 转发。家属指令到达老人端后，注入运行时，其后路径与本地指令一致，因此确认与"
      "动作组约束仍然有效。演示时可在局域网部署服务端，地址由配置指定。")

    H(doc, "3.11 主要时序（文字）", 2)
    P(doc, "（1）本地设置类指令：听写完成 → 路由命中端侧意图 → 打开设置页 → TTS 提示完成。")
    P(doc, "（2）拨号：模板或意图命中 → 解析联系人 → 确认问句 → 用户肯定 → 拨号意图执行。")
    P(doc, "（3）复杂跨应用：路由未命中 → 编排器循环观察与规划 → 可能触发动作组 → 结束播报。")
    P(doc, "（4）远程：协助码配对 → 家属下发文本 → 老人端运行时提交 → 同（2）或（3）。")

    # ========== 第四部分 数据 ==========
    H(doc, "第四部分  数据存储设计说明", 1)

    H(doc, "4.1 说明", 2)
    P(doc, "本系统不以大型业务库为核心。Android 端主要使用 SharedPreferences 保存配置与少量状态；"
      "模型与词表置于 assets；协助服务端使用轻量 H2 数据库保存配对与会话元数据。"
      "因此本部分以“逻辑数据说明”代替传统库表大全，满足竞赛“如必要再附数据库说明书”的弹性要求。")

    H(doc, "4.2 客户端持久化内容", 2)
    table(doc, ["数据类别", "主要内容", "备注"], [
        ["接口配置", "大模型与 ASR 密钥、模型名等", "勿写入公开仓库"],
        ["唤醒配置", "唤醒词、阈值、相关开关", ""],
        ["会话与记忆", "待确认事项、简要历史", "体积宜小"],
        ["照护配置", "亲属别名与电话、紧急文案、家地址", "供模板与系统能力使用"],
        ["预设指令", "短语到任务的映射", ""],
        ["协助绑定", "角色与绑定状态", ""],
        ["显示偏好", "主题等", ""],
    ])

    H(doc, "4.3 运行期主要结构", 2)
    table(doc, ["结构", "要点"], [
        ["AgentAction", "动作名、目标文本、输入文本、提示语、结束与确认标志"],
        ["界面观测", "节点摘要或截图，供规划使用"],
        ["步骤记录", "已执行动作与结果，供循环与动作组对齐"],
        ["意图预测", "标签、概率、margin"],
        ["协助消息", "类型与载荷"],
    ])

    H(doc, "4.4 服务端数据", 2)
    P(doc, "协助服务端保存房间、令牌与会话相关元数据，服务于短时协同，不承担长期用户画像存储。"
      "具体表结构以实现为准，部署演示环境时可使用嵌入式 H2。")

    H(doc, "4.5 数据流摘要", 2)
    num(doc, [
        "音频：采集 → VAD/KWS → ASR → 文本；",
        "决策：文本 → 路由 → 本地动作或规划结果；",
        "控制：动作 → 守卫 → 执行 → 新观测；",
        "协同：家属 → 中继 → 老人运行时；",
        "上云：界面摘要 → 脱敏 → 大模型。",
    ])

    # ========== 第五部分 界面 ==========
    H(doc, "第五部分  软件界面设计说明", 1)

    H(doc, "5.1 设计取向", 2)
    P(doc, "界面以分区清晰、状态可见为要：用户应能知道“现在能不能说话、系统在做什么、是否在等我确认”。"
      "文案力求口语化，布局避免堆叠过多装饰，减少老年用户的认知负担。")

    H(doc, "5.2 信息结构", 2)
    table(doc, ["页面", "主要内容", "对应能力"], [
        ["助手", "对话与卡片、任务步骤、建议语句、输入与语音入口", "日常下达指令"],
        ["设置", "唤醒、接口密钥、相关开关", "首次配置与调试"],
        ["协作", "角色选择、协助码、远程指令入口", "家属协助"],
        ["关于", "版本、模型信息、简要记忆", "说明与查验"],
    ])

    H(doc, "5.3 系统层界面", 2)
    num(doc, [
        "悬浮窗：在其他应用上提供入口；",
        "前台服务通知：说明唤醒或悬浮能力正在运行；",
        "确认相关通知：在需要时提醒用户回答确认问题。",
    ])

    H(doc, "5.4 典型界面流程", 2)
    P(doc, "首次使用：安装 → 开启无障碍（含必要组件）→ 授予麦克风与悬浮窗等 → 填写密钥与照护信息 →"
      "用建议语句完成一次成功体验。日常使用：唤醒或点击说话 → 观察步骤 → 如遇确认则回答是或否。"
      "远程协助：老人端展示协助码 → 家属连接 → 下发指令。")
    P(doc, "正式提交材料中，建议另附实际截图页；本合订本给出结构说明，便于与演示视频、PPT 对照。")

    # ========== 第六部分 使用与安装 ==========
    H(doc, "第六部分  用户操作与安装说明", 1)

    H(doc, "6.1 运行环境", 2)
    table(doc, ["项目", "要求"], [
        ["系统", "Android 9.0 及以上真机"],
        ["网络", "云端识别、规划、地图与协助需要联网"],
        ["权限", "麦克风、无障碍、悬浮窗等，按功能启用"],
        ["可选", "通知读取、通讯录、定位、Joy 输入法"],
    ])

    H(doc, "6.2 安装步骤", 2)
    num(doc, [
        "使用 Android Studio 或 Gradle 构建并安装 Debug/Release 包；",
        "复制并填写 local.properties 中的密钥与协助服务地址（勿提交到公开仓库）；",
        "如需远程协助，先启动 assist-server，并保证手机可访问该服务；",
        "打开应用，按引导开启无障碍主服务及微信支持组件（若演示微信相关功能）；",
        "授予麦克风、通知、悬浮窗等权限；推荐启用 Joy 输入助手；",
        "在设置中确认语音识别与大模型配置，在照护配置中填写亲属与家地址等。",
    ])

    H(doc, "6.3 基本操作", 2)
    table(doc, ["操作", "方法"], [
        ["语音下达", "说出唤醒词，或点击麦克风后说话"],
        ["文字下达", "在输入框输入指令并发送"],
        ["快捷体验", "点击建议语句"],
        ["确认发送/拨号", "听取问句后，明确表示同意或拒绝"],
        ["打断播报", "在播报过程中再次说话（视实现开启情况）"],
        ["远程协助", "协作页生成/展示协助码，家属端输入连接"],
    ])

    H(doc, "6.4 推荐体验语句", 2)
    bullets(doc, [
        "打开 WiFi / 放大字体 / 几点了；",
        "导航回家 / 我要回家；",
        "打电话给（已配置的亲属称呼）；",
        "帮我读一下未读消息；",
        "打开健康码 / 打开付款码；",
        "紧急呼救 / 叫家人帮忙（须先配置）。",
    ])

    H(doc, "6.5 常见问题", 2)
    table(doc, ["现象", "可能原因", "处理建议"], [
        ["不能点击别人家应用", "未开无障碍或组件不全", "检查系统设置中的无障碍开关"],
        ["听不懂 / 无反应", "无麦克风权限或网络/密钥问题", "检查权限与 API 配置"],
        ["微信里找不到控件", "未开微信支持组件", "同时开启主服务与支持组件"],
        ["协助连不上", "服务未启动或地址不对", "检查 assist-server 与配置地址"],
        ["总是问确认", "正在发送或拨号", "属正常安全设计，按提示回答即可"],
    ])

    # ========== 第七部分 出错、安全、测试 ==========
    H(doc, "第七部分  出错处理、安全与测试要点", 1)

    H(doc, "7.1 出错处理", 2)
    table(doc, ["情形", "处理原则"], [
        ["权限或服务缺失", "明确提示前往开启，不静默失败"],
        ["网络或云端失败", "可本地完成的继续；否则提示重试"],
        ["找不到目标控件", "尝试滚动或更换策略；超过步数则结束并说明"],
        ["意图把握不足", "不自动执行，转入澄清或其他路径"],
        ["用户拒绝确认", "取消动作并说明已取消"],
        ["协助异常", "提示检查协助码、网络与服务状态"],
    ])

    H(doc, "7.2 安全与保密", 2)
    num(doc, [
        "高权限能力须用户知情开启；",
        "发送、拨号须确认；",
        "限制无意义重复操作；",
        "上云上下文宜脱敏；",
        "密钥仅存本机；",
        "远程协助建议在可信网络演示，正式环境应进一步加强传输保护；",
        "紧急功能不能替代公共急救体系。",
    ])

    H(doc, "7.3 测试要点", 2)
    table(doc, ["编号", "场景", "期望结果"], [
        ["T01", "打开 WiFi 或放大字体", "走本地路径并成功"],
        ["T02", "导航回家", "调起地图导航"],
        ["T03", "拨号确认—同意", "确认后发起拨号"],
        ["T04", "拨号确认—拒绝", "不拨号并提示取消"],
        ["T05", "建议语句“几点了”", "播报时间"],
        ["T06", "关闭无障碍后下指令", "有明确提示"],
        ["T07", "断网后本地设置类指令", "仍可能完成"],
        ["T08", "协助码远程下发简单指令", "老人端执行"],
        ["T09", "近义意图干扰句", "不因低间隔而误自动执行"],
        ["T10", "发送类任务", "发送前出现确认"],
    ])

    # ========== 结束 ==========
    H(doc, "结束语", 1)
    P(doc,
      "JoyForOld 尝试在手机端把“听懂一句话”推进到“办完一件事”。本书从需求、结构、关键实现、"
      "数据与使用等方面作了系统说明，冀能为评审与后续改进提供清楚依据。完善之处，敬请各位老师指正。")
    P(doc, "（全文完）", align="center", indent=False, before=18)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for i in range(1, 5):
        try:
            doc.styles[f"Heading {i}"].font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass

    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(2.5))

    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("The 21th Hunan Collegiate Programming Contest")
    font(hr, size=9, color=RGBColor(0x66, 0x66, 0x66))

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("注意：文档内容要自动目录，双面打印，白色纸质胶装，一式2份；打印时请删除本页脚。")
    font(fr, size=8, cn="宋体", color=RGBColor(0x99, 0x99, 0x99))

    cover(doc)
    toc_page(doc)
    body(doc)

    saved = []
    for path in OUTS:
        try:
            doc.save(str(path))
            saved.append(path)
            print("Saved", path, path.stat().st_size)
        except PermissionError:
            print("Locked:", path.name)
        except Exception as e:
            print("Fail", path.name, e)
    if not saved:
        alt = DOCS / "JoyForOld-System-Manual-Complete.docx"
        doc.save(str(alt))
        print("Saved", alt)


if __name__ == "__main__":
    main()
